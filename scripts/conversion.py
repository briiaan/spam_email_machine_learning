import os
import pandas as pd
import re
from docx import Document
from bs4 import BeautifulSoup

# 📂 **Set the folder path**
folder_path = r"C:\Users\eric_\Desktop\Senior Project\cluster folder\datasets"
output_folder = os.path.join(folder_path, "docx_output")
os.makedirs(output_folder, exist_ok=True)

# 📌 **Read all CSV files in the folder**
file_types = ['.csv']
all_files = [f for f in os.listdir(folder_path) if any(f.endswith(ext) for ext in file_types)]

df_list = []
for file in all_files:
    file_path = os.path.join(folder_path, file)
    try:
        df_list.append(pd.read_csv(file_path, dtype=str, low_memory=False))  # Fix dtype warning
    except Exception as e:
        print(f"❌ Error loading {file}: {e}")

# 📌 **Merge all files into a single DataFrame**
df = pd.concat(df_list, ignore_index=True) if df_list else pd.DataFrame()

if df.empty:
    print("❌ No valid data found after merging files. Exiting...")
    exit()

print(f"✅ Successfully loaded {len(all_files)} files. Total records: {df.shape[0]}")

# **Ensure the 'text' column exists**
if 'message' in df.columns:
    df.rename(columns={'message': 'text'}, inplace=True)

df.dropna(subset=['text'], inplace=True)  # Remove empty rows

# 📌 **Text Cleaning Function**
def clean_text(text):
    if not isinstance(text, str):
        return ""
    text = BeautifulSoup(text, "html.parser").get_text()  # Remove HTML tags
    text = re.sub(r"http\S+|www\S+", "", text)  # Remove URLs
    text = re.sub(r"[^a-zA-Z\s]", "", text)  # Remove special characters, numbers
    text = text.lower().strip()  # Convert to lowercase
    return text

df['cleaned_text'] = df['text'].apply(clean_text)

print(f"✅ Text cleaning complete. Total processed records: {df.shape[0]}")

# 📌 **Batch Processing - Save Emails in Groups**
batch_size = 5000  # Set batch size (adjust as needed)
batch_number = 1
doc = Document()

for index, row in df.iterrows():
    doc.add_heading(f"Email {index + 1}", level=1)
    doc.add_paragraph(row['cleaned_text'])

    # Save every batch of 'batch_size' emails
    if (index + 1) % batch_size == 0 or index == len(df) - 1:
        file_path = os.path.join(output_folder, f"batch_{batch_number}.docx")
        doc.save(file_path)
        print(f"✅ Saved batch {batch_number}: {file_path}")
        doc = Document()  # Reset for the next batch
        batch_number += 1

print(f"✅ Successfully saved {batch_number - 1} DOCX batch files in: {output_folder}")
